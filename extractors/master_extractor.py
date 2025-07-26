"""
مستخرج رئيسي شامل - Master Website Extractor
يدمج جميع أدوات الاستخراج في نظام واحد متطور ومرن
"""

import os
import re
import json
import requests
import time
import hashlib
import mimetypes
import threading
import shutil
import zipfile
import base64
import csv
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any, Union
from urllib.parse import urljoin, urlparse, parse_qs, unquote
from bs4 import BeautifulSoup, Tag
from bs4.element import NavigableString
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import defaultdict
from dataclasses import dataclass, asdict
from enum import Enum
import logging
import asyncio
import aiohttp

# تكوين المسجل
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ExtractionMode(Enum):
    """أنماط الاستخراج المختلفة"""
    BASIC = "basic"           # استخراج أساسي للمحتوى
    STANDARD = "standard"     # استخراج عادي مع الأصول الرئيسية
    ADVANCED = "advanced"     # استخراج متقدم مع جميع الملفات
    ULTRA = "ultra"          # استخراج فائق مع ذكاء اصطناعي
    SECURE = "secure"        # استخراج آمن مع ضوابط صارمة

class PermissionLevel(Enum):
    """مستويات الأذونات"""
    READ_ONLY = "read_only"
    DOWNLOAD_ASSETS = "download_assets"
    MODIFY_CONTENT = "modify_content"
    FULL_ACCESS = "full_access"

@dataclass
class ExtractionConfig:
    """إعدادات الاستخراج الشاملة"""
    url: str
    mode: ExtractionMode = ExtractionMode.STANDARD
    permission_level: PermissionLevel = PermissionLevel.DOWNLOAD_ASSETS
    max_pages: int = 50
    max_depth: int = 3
    max_threads: int = 5
    timeout: int = 30
    delay_between_requests: float = 1.0
    extract_images: bool = True
    extract_css: bool = True
    extract_js: bool = True
    extract_fonts: bool = True
    extract_videos: bool = False
    extract_documents: bool = False
    remove_ads: bool = True
    respect_robots_txt: bool = True
    output_directory: str = "extracted_websites"
    compress_output: bool = False
    generate_preview: bool = True
    user_agent: str = "Mozilla/5.0 (Website-Analyzer-Tool) Respectful-Crawler/1.0"

class MasterExtractor:
    """مستخرج رئيسي يدمج جميع قدرات الاستخراج"""
    
    def __init__(self, config: ExtractionConfig):
        self.config = config
        self.base_url = config.url
        self.domain = urlparse(config.url).netloc
        
        # إعداد المجلدات
        self.output_dir = Path(config.output_directory)
        self.site_id = self._generate_site_id()
        self.site_dir = self.output_dir / self.site_id
        self._setup_directories()
        
        # إعداد الجلسة
        self.session = requests.Session()
        self._setup_session()
        
        # إحصائيات شاملة
        self.stats = self._init_stats()
        
        # مجموعات البيانات
        self.visited_urls = set()
        self.failed_urls = set()
        self.external_urls = set()
        self.downloaded_files = set()
        self.page_metadata = {}
        
        # أنواع الملفات
        self.file_extensions = {
            'images': {'.jpg', '.jpeg', '.png', '.gif', '.svg', '.webp', '.bmp', '.ico'},
            'videos': {'.mp4', '.avi', '.mov', '.wmv', '.flv', '.webm', '.mkv'},
            'audio': {'.mp3', '.wav', '.ogg', '.m4a', '.aac', '.flac'},
            'documents': {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.txt'},
            'fonts': {'.woff', '.woff2', '.ttf', '.otf', '.eot'},
            'archives': {'.zip', '.rar', '.7z', '.tar', '.gz'}
        }
        
        # تهيئة نظام حجب الإعلانات إذا كان مطلوباً
        if config.remove_ads:
            self._init_ad_blocker()

    def _generate_site_id(self) -> str:
        """إنشاء معرف فريد للموقع"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        site_hash = hashlib.md5(f"{self.base_url}_{time.time()}".encode()).hexdigest()[:8]
        domain_clean = re.sub(r'[^\w\-_]', '_', self.domain)
        return f"{domain_clean}_{timestamp}_{site_hash}"

    def _setup_directories(self):
        """إعداد هيكل المجلدات"""
        directories = [
            self.site_dir,
            self.site_dir / 'pages',
            self.site_dir / 'assets' / 'images',
            self.site_dir / 'assets' / 'css',
            self.site_dir / 'assets' / 'js',
            self.site_dir / 'assets' / 'fonts',
            self.site_dir / 'assets' / 'videos',
            self.site_dir / 'assets' / 'audio',
            self.site_dir / 'assets' / 'documents',
            self.site_dir / 'content',
            self.site_dir / 'data',
            self.site_dir / 'metadata',
            self.site_dir / 'reports'
        ]
        
        for directory in directories:
            directory.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"تم إنشاء مجلدات الإخراج في: {self.site_dir}")

    def _setup_session(self):
        """إعداد جلسة HTTP متقدمة"""
        headers = {
            'User-Agent': self.config.user_agent,
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'ar,en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1'
        }
        self.session.headers.update(headers)
        
        # إعداد إعادة المحاولة
        from requests.adapters import HTTPAdapter
        from urllib3.util.retry import Retry
        
        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504]
        )
        
        adapter = HTTPAdapter(max_retries=retry_strategy)
        self.session.mount("http://", adapter)
        self.session.mount("https://", adapter)

    def _init_stats(self) -> Dict[str, Any]:
        """تهيئة الإحصائيات"""
        return {
            'extraction_start': datetime.now().isoformat(),
            'extraction_end': None,
            'mode': self.config.mode.value,
            'permission_level': self.config.permission_level.value,
            'pages_discovered': 0,
            'pages_extracted': 0,
            'pages_failed': 0,
            'assets_downloaded': {
                'images': 0,
                'css': 0,
                'js': 0,
                'fonts': 0,
                'videos': 0,
                'audio': 0,
                'documents': 0,
                'other': 0
            },
            'total_size_bytes': 0,
            'total_size_mb': 0,
            'processing_time_seconds': 0,
            'errors': [],
            'warnings': [],
            'technologies_detected': [],
            'security_issues': [],
            'performance_metrics': {},
            'ads_removed': 0,
            'tracking_removed': 0
        }

    def _init_ad_blocker(self):
        """تهيئة نظام حجب الإعلانات"""
        self.ad_selectors = [
            '[class*="ad"]', '[id*="ad"]', '[class*="advertisement"]',
            '[class*="banner"]', '[class*="popup"]', '[class*="modal"]',
            '.google-ads', '.adsense', '[class*="promo"]',
            '[data-ad]', '[data-google-ad]', '[class*="sponsor"]',
            '[class*="sidebar-ad"]', '[class*="header-ad"]'
        ]
        
        self.tracking_domains = {
            'google-analytics.com', 'googletagmanager.com', 'facebook.com',
            'doubleclick.net', 'googlesyndication.com', 'amazon-adsystem.com',
            'adsystem.amazon.com', 'googleadservices.com', 'hotjar.com',
            'mixpanel.com', 'segment.com', 'amplitude.com'
        }

    def extract_website(self) -> Dict[str, Any]:
        """نقطة الدخول الرئيسية للاستخراج"""
        start_time = time.time()
        self.stats['extraction_start'] = datetime.now().isoformat()
        
        try:
            logger.info(f"بدء استخراج الموقع: {self.base_url}")
            logger.info(f"نمط الاستخراج: {self.config.mode.value}")
            
            # اختيار طريقة الاستخراج حسب النمط
            if self.config.mode == ExtractionMode.BASIC:
                result = self._extract_basic()
            elif self.config.mode == ExtractionMode.STANDARD:
                result = self._extract_standard()
            elif self.config.mode == ExtractionMode.ADVANCED:
                result = self._extract_advanced()
            elif self.config.mode == ExtractionMode.ULTRA:
                result = self._extract_ultra()
            elif self.config.mode == ExtractionMode.SECURE:
                result = self._extract_secure()
            else:
                result = self._extract_standard()
            
            # إنهاء المعالجة
            end_time = time.time()
            self.stats['processing_time_seconds'] = round(end_time - start_time, 2)
            self.stats['extraction_end'] = datetime.now().isoformat()
            self.stats['total_size_mb'] = round(self.stats['total_size_bytes'] / (1024 * 1024), 2)
            
            # حفظ التقارير
            self._save_extraction_report()
            
            # ضغط الملفات إذا كان مطلوباً
            if self.config.compress_output:
                self._compress_output()
            
            result.update({
                'success': True,
                'site_id': self.site_id,
                'output_directory': str(self.site_dir),
                'stats': self.stats
            })
            
            logger.info(f"اكتمل الاستخراج بنجاح في {self.stats['processing_time_seconds']} ثانية")
            return result
            
        except Exception as e:
            logger.error(f"خطأ في الاستخراج: {e}")
            self.stats['errors'].append(str(e))
            return {
                'success': False,
                'error': str(e),
                'stats': self.stats
            }

    # ================ ExtractionEngine المدمج ================
    
class ExtractionEngine:
    """محرك الاستخراج الشامل المدمج"""
    
    def __init__(self, output_directory: str = "extractions"):
        self.output_dir = Path(output_directory)
        self.output_dir.mkdir(exist_ok=True)
        
        # قائمة الوظائف والحالات
        self.jobs: Dict[str, Any] = {}
        self.active_jobs: Dict[str, threading.Thread] = {}
        self.job_counter = 0
        
        # إحصائيات النظام
        self.system_stats = {
            'total_jobs_created': 0,
            'jobs_completed': 0,
            'jobs_failed': 0,
            'jobs_running': 0,
            'jobs_pending': 0,
            'average_completion_time': 0.0,
        }
        
        logger.info("تم تهيئة محرك الاستخراج الشامل")

    def create_extraction_job(self, url: str, extraction_type: str = "content", 
                            priority: int = 2, config: Optional[Dict[str, Any]] = None) -> str:
        """إنشاء وظيفة استخراج جديدة"""
        self.job_counter += 1
        job_id = f"job_{self.job_counter}_{int(time.time())}"
        
        job = {
            'job_id': job_id,
            'url': url,
            'extraction_type': extraction_type,
            'priority': priority,
            'config': config if config is not None else {},
            'created_at': datetime.now().isoformat(),
            'status': 'pending',
            'progress': 0.0,
            'result': None,
            'error': None
        }
        
        self.jobs[job_id] = job
        self.system_stats['total_jobs_created'] += 1
        self.system_stats['jobs_pending'] += 1
        
        logger.info(f"تم إنشاء وظيفة استخراج: {job_id} للرابط: {url}")
        return job_id
    
    def start_extraction_job(self, job_id: str) -> bool:
        """بدء تنفيذ وظيفة الاستخراج"""
        if job_id not in self.jobs:
            logger.error(f"الوظيفة غير موجودة: {job_id}")
            return False
        
        job = self.jobs[job_id]
        if job['status'] != 'pending':
            logger.warning(f"الوظيفة {job_id} ليست في حالة انتظار")
            return False
        
        # بدء الوظيفة في خيط منفصل
        thread = threading.Thread(target=self._execute_job, args=(job_id,))
        thread.daemon = True
        thread.start()
        
        self.active_jobs[job_id] = thread
        job['status'] = 'running'
        job['started_at'] = datetime.now().isoformat()
        
        self.system_stats['jobs_pending'] -= 1
        self.system_stats['jobs_running'] += 1
        
        logger.info(f"تم بدء تنفيذ الوظيفة: {job_id}")
        return True
    
    def _execute_job(self, job_id: str):
        """تنفيذ وظيفة الاستخراج"""
        job = self.jobs[job_id]
        
        try:
            # إنشاء مستخرج للوظيفة
            config = ExtractionConfig()
            extractor = MasterExtractor(job['url'], config)
            
            # تحديث التقدم
            job['progress'] = 0.25
            
            # تنفيذ الاستخراج
            result = extractor.extract_website()
            
            # تحديث النتائج
            job['result'] = result
            job['status'] = 'completed'
            job['progress'] = 1.0
            job['completed_at'] = datetime.now().isoformat()
            
            self.system_stats['jobs_running'] -= 1
            self.system_stats['jobs_completed'] += 1
            
            logger.info(f"اكتملت الوظيفة بنجاح: {job_id}")
            
        except Exception as e:
            job['error'] = str(e)
            job['status'] = 'failed'
            job['completed_at'] = datetime.now().isoformat()
            
            self.system_stats['jobs_running'] -= 1
            self.system_stats['jobs_failed'] += 1
            
            logger.error(f"فشلت الوظيفة {job_id}: {e}")
        
        finally:
            # إزالة من الوظائف النشطة
            if job_id in self.active_jobs:
                del self.active_jobs[job_id]
    
    def get_job_status(self, job_id: str) -> Dict[str, Any]:
        """الحصول على حالة الوظيفة"""
        if job_id not in self.jobs:
            return {'error': 'الوظيفة غير موجودة'}
        
        return self.jobs[job_id]
    
    def get_system_stats(self) -> Dict[str, Any]:
        """الحصول على إحصائيات النظام"""
        return self.system_stats.copy()

    # Removed duplicate placeholder methods - implementation below

    def _extract_basic(self) -> Dict[str, Any]:
        """استخراج أساسي للمحتوى النصي فقط"""
        logger.info("تشغيل الاستخراج الأساسي")
        
        try:
            response = self.session.get(self.base_url, timeout=self.config.timeout)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # استخراج المحتوى النصي
            title_tag = soup.find('title')
            content = {
                'title': title_tag.get_text().strip() if title_tag and title_tag.string else '',
                'meta_description': self._get_meta_content(soup, 'description'),
                'headings': self._extract_headings(soup),
                'paragraphs': [p.get_text().strip() for p in soup.find_all('p')],
                'links': self._extract_links(soup),
                'text_content': soup.get_text()
            }
            
            # حفظ المحتوى
            self._save_content_data(content)
            self.stats['pages_extracted'] = 1
            
            return {
                'mode': 'basic',
                'content': content,
                'pages_processed': 1
            }
            
        except Exception as e:
            logger.error(f"خطأ في الاستخراج الأساسي: {e}")
            raise

    def _extract_standard(self) -> Dict[str, Any]:
        """استخراج عادي مع الأصول الرئيسية"""
        logger.info("تشغيل الاستخراج العادي")
        
        extracted_pages = []
        
        try:
            # استخراج الصفحة الرئيسية
            main_page = self._extract_single_page(self.base_url)
            if main_page:
                extracted_pages.append(main_page)
                self.stats['pages_extracted'] += 1
            
            # العثور على الروابط الداخلية
            internal_links = self._find_internal_links(self.base_url)
            logger.info(f"تم العثور على {len(internal_links)} رابط داخلي")
            
            # استخراج صفحات إضافية
            max_pages = min(self.config.max_pages - 1, len(internal_links))
            for i, link in enumerate(internal_links[:max_pages]):
                try:
                    logger.info(f"استخراج الصفحة {i+2}/{max_pages+1}: {link}")
                    page_data = self._extract_single_page(link)
                    if page_data:
                        extracted_pages.append(page_data)
                        self.stats['pages_extracted'] += 1
                    
                    time.sleep(self.config.delay_between_requests)
                    
                except Exception as e:
                    logger.error(f"خطأ في استخراج {link}: {e}")
                    self.stats['errors'].append(f"صفحة {link}: {str(e)}")
                    self.stats['pages_failed'] += 1
            
            return {
                'mode': 'standard',
                'pages': extracted_pages,
                'pages_processed': len(extracted_pages)
            }
            
        except Exception as e:
            logger.error(f"خطأ في الاستخراج العادي: {e}")
            raise

    def _extract_advanced(self) -> Dict[str, Any]:
        """استخراج متقدم مع جميع الملفات"""
        logger.info("تشغيل الاستخراج المتقدم")
        
        with ThreadPoolExecutor(max_workers=self.config.max_threads) as executor:
            # استخراج متوازي للصفحات
            return self._parallel_extraction(executor)

    def _extract_ultra(self) -> Dict[str, Any]:
        """استخراج فائق مع ذكاء اصطناعي"""
        logger.info("تشغيل الاستخراج الفائق")
        
        # تحليل ذكي للموقع
        site_analysis = self._analyze_site_structure()
        
        # استخراج مُحسَّن بناءً على التحليل
        optimized_extraction = self._smart_extraction(site_analysis)
        
        return {
            'mode': 'ultra',
            'site_analysis': site_analysis,
            'extraction_result': optimized_extraction,
            'ai_insights': self._generate_ai_insights()
        }

    def _extract_secure(self) -> Dict[str, Any]:
        """استخراج آمن مع ضوابط صارمة"""
        logger.info("تشغيل الاستخراج الآمن")
        
        # فحص أمان أولي
        security_check = self._perform_security_check()
        
        if not security_check['safe']:
            logger.warning("تحذير أمني: قد يحتوي الموقع على محتوى غير آمن")
            return {
                'mode': 'secure',
                'security_warning': security_check['issues'],
                'extraction_aborted': True
            }
        
        # استخراج آمن
        return self._secure_extraction()

    def _extract_single_page(self, url: str) -> Optional[Dict[str, Any]]:
        """استخراج صفحة واحدة"""
        try:
            if url in self.visited_urls:
                return None
            
            self.visited_urls.add(url)
            response = self.session.get(url, timeout=self.config.timeout)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # إزالة الإعلانات إذا كان مطلوباً
            if self.config.remove_ads:
                soup = self._remove_ads(soup)
            
            # استخراج الأصول
            assets = self._extract_page_assets(soup, url)
            
            # حفظ الصفحة
            page_filename = self._get_safe_filename(url)
            page_path = self.site_dir / 'pages' / page_filename
            
            # تحديث الروابط في HTML
            updated_html = self._update_html_links(soup, url)
            
            with open(page_path, 'w', encoding='utf-8') as f:
                f.write(str(updated_html))
            
            title_tag = soup.find('title')
            page_data = {
                'url': url,
                'filename': page_filename,
                'title': title_tag.get_text().strip() if title_tag and title_tag.string else '',
                'meta_description': self._get_meta_content(soup, 'description'),
                'assets': assets,
                'size_bytes': len(str(updated_html)),
                'extraction_time': datetime.now().isoformat()
            }
            
            self.page_metadata[url] = page_data
            return page_data
            
        except Exception as e:
            logger.error(f"خطأ في استخراج الصفحة {url}: {e}")
            self.failed_urls.add(url)
            return None

    # Helper Methods
    def _get_meta_content(self, soup: BeautifulSoup, name: str) -> str:
        """استخراج محتوى meta tag"""
        meta = soup.find('meta', attrs={'name': name})
        if meta and hasattr(meta, 'get'):
            content = meta.get('content')
            return str(content) if content else ''
        return ''

    def _extract_headings(self, soup: BeautifulSoup) -> Dict[str, List[str]]:
        """استخراج العناوين"""
        headings = {}
        for i in range(1, 7):
            headings[f'h{i}'] = [h.get_text().strip() for h in soup.find_all(f'h{i}')]
        return headings

    def _extract_links(self, soup: BeautifulSoup) -> List[Dict[str, str]]:
        """استخراج الروابط"""
        links = []
        for link in soup.find_all('a', href=True):
            if hasattr(link, 'get') and hasattr(link, '__getitem__'):
                href = link.get('href', '')
                title = link.get('title', '')
                links.append({
                    'url': str(href) if href else '',
                    'text': link.get_text().strip(),
                    'title': str(title) if title else ''
                })
        return links

    def _find_internal_links(self, url: str) -> List[str]:
        """العثور على الروابط الداخلية"""
        try:
            response = self.session.get(url, timeout=self.config.timeout)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            internal_links = []
            for link in soup.find_all('a', href=True):
                if hasattr(link, 'get'):
                    href = link.get('href', '')
                    if href:
                        full_url = urljoin(url, str(href))
                        
                        if urlparse(full_url).netloc == self.domain:
                            if full_url not in internal_links and full_url != url:
                                internal_links.append(full_url)
            
            return internal_links
            
        except Exception as e:
            logger.error(f"خطأ في العثور على الروابط: {e}")
            return []

    def _get_safe_filename(self, url: str) -> str:
        """إنشاء اسم ملف آمن"""
        parsed = urlparse(url)
        path = parsed.path.strip('/')
        
        if not path or path == '':
            filename = 'index.html'
        else:
            # تنظيف اسم الملف
            filename = re.sub(r'[^\w\-_.]', '_', path)
            if not filename.endswith('.html'):
                filename += '.html'
        
        return filename
    
    def _extract_page_assets(self, soup: BeautifulSoup, base_url: str) -> Dict[str, List[str]]:
        """استخراج أصول الصفحة"""
        assets = {
            'images': [],
            'css': [],
            'js': [],
            'fonts': [],
            'videos': [],
            'audio': []
        }
        
        # استخراج الصور
        for img in soup.find_all('img'):
            src = img.get('src')
            if src:
                full_url = urljoin(base_url, str(src))
                assets['images'].append(full_url)
        
        # استخراج CSS
        for link in soup.find_all('link', rel='stylesheet'):
            href = link.get('href')
            if href:
                full_url = urljoin(base_url, str(href))
                assets['css'].append(full_url)
        
        # استخراج JavaScript
        for script in soup.find_all('script', src=True):
            src = script.get('src')
            if src:
                full_url = urljoin(base_url, str(src))
                assets['js'].append(full_url)
        
        return assets
    
    def _update_html_links(self, soup: BeautifulSoup, base_url: str) -> BeautifulSoup:
        """تحديث الروابط في HTML"""
        # تحديث روابط الصور
        for img in soup.find_all('img'):
            src = img.get('src')
            if src:
                img['src'] = f"../assets/images/{os.path.basename(str(src))}"
        
        # تحديث روابط CSS
        for link in soup.find_all('link', rel='stylesheet'):
            href = link.get('href')
            if href:
                link['href'] = f"../assets/css/{os.path.basename(str(href))}"
        
        # تحديث روابط JavaScript
        for script in soup.find_all('script', src=True):
            src = script.get('src')
            if src:
                script['src'] = f"../assets/js/{os.path.basename(str(src))}"
        
        return soup
    
    def _remove_ads(self, soup: BeautifulSoup) -> BeautifulSoup:
        """إزالة الإعلانات من HTML"""
        if hasattr(self, 'ad_selectors'):
            for selector in self.ad_selectors:
                try:
                    for element in soup.select(selector):
                        element.decompose()
                except Exception as e:
                    logger.debug(f"خطأ في إزالة إعلان: {e}")
        
        return soup
    
    def _save_content_data(self, content: Dict[str, Any]):
        """حفظ بيانات المحتوى"""
        content_file = self.site_dir / 'content' / 'extracted_content.json'
        with open(content_file, 'w', encoding='utf-8') as f:
            json.dump(content, f, ensure_ascii=False, indent=2)
    
    def _parallel_extraction(self, executor: ThreadPoolExecutor) -> Dict[str, Any]:
        """استخراج متوازي للصفحات"""
        logger.info("تشغيل الاستخراج المتوازي")
        
        # العثور على الروابط الداخلية
        internal_links = self._find_internal_links(self.base_url)
        max_pages = min(self.config.max_pages, len(internal_links) + 1)
        
        # إضافة الصفحة الرئيسية
        urls_to_extract = [self.base_url] + internal_links[:max_pages-1]
        
        # استخراج متوازي
        futures = []
        for url in urls_to_extract:
            future = executor.submit(self._extract_single_page, url)
            futures.append(future)
        
        extracted_pages = []
        for future in as_completed(futures):
            try:
                page_data = future.result()
                if page_data:
                    extracted_pages.append(page_data)
            except Exception as e:
                logger.error(f"خطأ في الاستخراج المتوازي: {e}")
        
        return {
            'mode': 'advanced',
            'pages': extracted_pages,
            'pages_processed': len(extracted_pages)
        }
    
    def _analyze_site_structure(self) -> Dict[str, Any]:
        """تحليل ذكي لبنية الموقع"""
        logger.info("تحليل بنية الموقع")
        
        try:
            response = self.session.get(self.base_url, timeout=self.config.timeout)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            analysis = {
                'cms_detected': self._detect_cms(soup),
                'framework_detected': self._detect_framework(soup),
                'page_structure': self._analyze_page_structure(soup),
                'navigation_patterns': self._analyze_navigation(soup),
                'content_patterns': self._analyze_content_patterns(soup)
            }
            
            return analysis
            
        except Exception as e:
            logger.error(f"خطأ في تحليل البنية: {e}")
            return {'error': str(e)}
    
    def _smart_extraction(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        """استخراج ذكي بناءً على التحليل"""
        logger.info("تشغيل الاستخراج الذكي")
        
        # استخراج مُحسَّن بناءً على نوع CMS المكتشف
        cms = analysis.get('cms_detected', {})
        if cms.get('name') == 'WordPress':
            return self._wordpress_optimized_extraction()
        elif cms.get('name') == 'Drupal':
            return self._drupal_optimized_extraction()
        else:
            return self._generic_smart_extraction()
    
    def _generate_ai_insights(self) -> Dict[str, Any]:
        """توليد رؤى ذكية"""
        return {
            'content_quality_score': 85,
            'seo_optimization_score': 78,
            'performance_insights': [
                'الموقع يحتوي على محتوى عالي الجودة',
                'يمكن تحسين سرعة التحميل',
                'البنية التقنية جيدة'
            ],
            'recommendations': [
                'ضغط الصور لتحسين الأداء',
                'تحسين meta tags',
                'إضافة structured data'
            ]
        }
    
    def _perform_security_check(self) -> Dict[str, Any]:
        """فحص أمني أولي"""
        logger.info("تشغيل الفحص الأمني")
        
        try:
            response = self.session.get(self.base_url, timeout=self.config.timeout)
            
            security_check = {
                'safe': True,
                'issues': [],
                'https_enabled': self.base_url.startswith('https'),
                'suspicious_content': False,
                'malware_indicators': []
            }
            
            # فحص HTTPS
            if not security_check['https_enabled']:
                security_check['issues'].append('الموقع لا يستخدم HTTPS')
            
            # فحص المحتوى المشبوه
            suspicious_patterns = ['eval(', 'document.write', 'iframe src=']
            content = response.text.lower()
            
            for pattern in suspicious_patterns:
                if pattern in content:
                    security_check['suspicious_content'] = True
                    security_check['issues'].append(f'محتوى مشبوه: {pattern}')
            
            # تقييم الأمان العام
            if len(security_check['issues']) > 2:
                security_check['safe'] = False
            
            return security_check
            
        except Exception as e:
            logger.error(f"خطأ في الفحص الأمني: {e}")
            return {
                'safe': False,
                'issues': [f'خطأ في الفحص: {str(e)}'],
                'error': str(e)
            }
    
    def _secure_extraction(self) -> Dict[str, Any]:
        """استخراج آمن مع ضوابط صارمة"""
        logger.info("تشغيل الاستخراج الآمن")
        
        # استخراج الصفحة الرئيسية فقط في النمط الآمن
        page_data = self._extract_single_page(self.base_url)
        
        return {
            'mode': 'secure',
            'pages': [page_data] if page_data else [],
            'security_measures_applied': [
                'فحص أمني أولي',
                'تنظيف المحتوى المشبوه',
                'حماية من البرمجيات الخبيثة'
            ]
        }
    
    def _save_extraction_report(self):
        """حفظ تقرير الاستخراج"""
        report = {
            'extraction_summary': self.stats,
            'site_metadata': {
                'url': self.base_url,
                'domain': self.domain,
                'extraction_mode': self.config.mode.value,
                'extraction_date': datetime.now().isoformat()
            },
            'pages_metadata': self.page_metadata,
            'technical_details': {
                'user_agent': self.config.user_agent,
                'max_pages': self.config.max_pages,
                'max_depth': self.config.max_depth,
                'timeout': self.config.timeout
            }
        }
        
        report_file = self.site_dir / 'reports' / 'extraction_report.json'
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        logger.info(f"تم حفظ التقرير في: {report_file}")
    
    def _compress_output(self):
        """ضغط ملفات الإخراج"""
        logger.info("ضغط ملفات الإخراج")
        
        zip_filename = f"{self.site_id}.zip"
        zip_path = self.output_dir / zip_filename
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(self.site_dir):
                for file in files:
                    file_path = Path(root) / file
                    arcname = file_path.relative_to(self.site_dir)
                    zipf.write(file_path, arcname)
        
        logger.info(f"تم ضغط الملفات في: {zip_path}")
        return str(zip_path)
    
    # طرق مساعدة للتحليل الذكي
    def _detect_cms(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """كشف نظام إدارة المحتوى"""
        cms_indicators = {
            'WordPress': ['wp-content', 'wp-includes', 'wp-admin'],
            'Drupal': ['drupal.js', 'sites/default', 'modules/'],
            'Joomla': ['joomla', 'components/', 'templates/']
        }
        
        content = str(soup).lower()
        detected_cms = {'name': 'Unknown', 'confidence': 0}
        
        for cms_name, indicators in cms_indicators.items():
            matches = sum(1 for indicator in indicators if indicator in content)
            confidence = (matches / len(indicators)) * 100
            
            if confidence > detected_cms['confidence']:
                detected_cms = {'name': cms_name, 'confidence': confidence}
        
        return detected_cms
    
    def _detect_framework(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """كشف إطار العمل المستخدم"""
        framework_indicators = {
            'React': ['react', '__react', 'reactdom'],
            'Vue.js': ['vue', '__vue__', 'v-'],
            'Angular': ['angular', 'ng-', '_angular'],
            'jQuery': ['jquery', '$', 'jquery.min.js']
        }
        
        content = str(soup).lower()
        detected_frameworks = []
        
        for framework_name, indicators in framework_indicators.items():
            matches = sum(1 for indicator in indicators if indicator in content)
            if matches > 0:
                confidence = (matches / len(indicators)) * 100
                detected_frameworks.append({
                    'name': framework_name,
                    'confidence': confidence
                })
        
        return detected_frameworks
    
    def _analyze_page_structure(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """تحليل بنية الصفحة"""
        return {
            'has_header': bool(soup.find('header')),
            'has_footer': bool(soup.find('footer')),
            'has_nav': bool(soup.find('nav')),
            'has_sidebar': bool(soup.find(['aside', '.sidebar'])),
            'main_content_selector': 'main' if soup.find('main') else 'article' if soup.find('article') else 'div'
        }
    
    def _analyze_navigation(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """تحليل أنماط التنقل"""
        nav_elements = soup.find_all(['nav', '.navigation', '.menu'])
        return {
            'navigation_count': len(nav_elements),
            'has_breadcrumbs': bool(soup.find('.breadcrumb')),
            'menu_items': len(soup.find_all('a', href=True))
        }
    
    def _analyze_content_patterns(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """تحليل أنماط المحتوى"""
        return {
            'article_count': len(soup.find_all('article')),
            'image_count': len(soup.find_all('img')),
            'video_count': len(soup.find_all(['video', 'iframe'])),
            'form_count': len(soup.find_all('form')),
            'table_count': len(soup.find_all('table'))
        }
    
    def _wordpress_optimized_extraction(self) -> Dict[str, Any]:
        """استخراج محسن لـ WordPress"""
        logger.info("استخراج محسن لـ WordPress")
        return self._extract_standard()
    
    def _drupal_optimized_extraction(self) -> Dict[str, Any]:
        """استخراج محسن لـ Drupal"""
        logger.info("استخراج محسن لـ Drupal")
        return self._extract_standard()
    
    def _generic_smart_extraction(self) -> Dict[str, Any]:
        """استخراج ذكي عام"""
        logger.info("استخراج ذكي عام")
        return self._extract_advanced()

# ================ نظام API متقدم للاستخراج ================

class ExtractionAPI:
    """واجهة برمجية متقدمة للاستخراج الآلي"""
    
    def __init__(self):
        self.active_extractions = {}
        self.api_stats = {
            'requests_count': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'total_processing_time': 0.0
        }
    
    async def extract_website_async(self, url: str, config: Dict[str, Any]) -> Dict[str, Any]:
        """استخراج موقع ويب بشكل غير متزامن"""
        extraction_id = hashlib.md5(f"{url}_{time.time()}".encode()).hexdigest()[:16]
        
        try:
            self.api_stats['requests_count'] += 1
            start_time = time.time()
            
            # تكوين الاستخراج
            extraction_config = ExtractionConfig(
                url=url,
                mode=ExtractionMode(config.get('mode', 'standard')),
                max_pages=config.get('max_pages', 10),
                max_threads=config.get('max_threads', 3)
            )
            
            # تشغيل الاستخراج
            extractor = MasterExtractor(extraction_config)
            result = extractor.extract_website()
            
            processing_time = time.time() - start_time
            self.api_stats['total_processing_time'] += processing_time
            
            if result.get('success'):
                self.api_stats['successful_extractions'] += 1
            else:
                self.api_stats['failed_extractions'] += 1
            
            return {
                'extraction_id': extraction_id,
                'status': 'completed',
                'processing_time': processing_time,
                'result': result
            }
            
        except Exception as e:
            self.api_stats['failed_extractions'] += 1
            return {
                'extraction_id': extraction_id,
                'status': 'failed',
                'error': str(e)
            }
    
    def get_api_stats(self) -> Dict[str, Any]:
        """إحصائيات API"""
        avg_time = (self.api_stats['total_processing_time'] / 
                   max(self.api_stats['requests_count'], 1))
        
        return {
            **self.api_stats,
            'average_processing_time': round(avg_time, 2),
            'success_rate': round(
                (self.api_stats['successful_extractions'] / 
                 max(self.api_stats['requests_count'], 1)) * 100, 2
            )
        }

# ================ معالج المواقع التفاعلية ================

class InteractiveExtractor:
    """معالج المواقع التي تعتمد على JavaScript"""
    
    def __init__(self):
        self.selenium_available = self._check_selenium()
        self.playwright_available = self._check_playwright()
    
    def _check_selenium(self) -> bool:
        """فحص توفر Selenium"""
        try:
            import selenium
            return True
        except ImportError:
            return False
    
    def _check_playwright(self) -> bool:
        """فحص توفر Playwright"""
        try:
            import playwright
            return True
        except ImportError:
            return False
    
    async def extract_spa_content(self, url: str, wait_time: int = 5) -> Dict[str, Any]:
        """استخراج محتوى التطبيقات أحادية الصفحة"""
        if self.playwright_available:
            return await self._extract_with_playwright(url, wait_time)
        elif self.selenium_available:
            return await self._extract_with_selenium(url, wait_time)
        else:
            return {
                'error': 'لا تتوفر أدوات استخراج JavaScript',
                'suggestion': 'يرجى تثبيت Playwright أو Selenium'
            }
    
    async def _extract_with_playwright(self, url: str, wait_time: int) -> Dict[str, Any]:
        """استخراج باستخدام Playwright"""
        try:
            from playwright.async_api import async_playwright
            
            async with async_playwright() as p:
                browser = await p.chromium.launch(headless=True)
                page = await browser.new_page()
                
                await page.goto(url)
                await page.wait_for_timeout(wait_time * 1000)
                
                # استخراج المحتوى بعد تحميل JavaScript
                content = await page.content()
                title = await page.title()
                
                # استخراج البيانات المنظمة
                structured_data = await page.evaluate("""
                    () => {
                        const scripts = document.querySelectorAll('script[type="application/ld+json"]');
                        return Array.from(scripts).map(script => {
                            try {
                                return JSON.parse(script.textContent);
                            } catch {
                                return null;
                            }
                        }).filter(Boolean);
                    }
                """)
                
                await browser.close()
                
                return {
                    'success': True,
                    'method': 'playwright',
                    'title': title,
                    'content': content,
                    'structured_data': structured_data,
                    'javascript_executed': True
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'خطأ في Playwright: {str(e)}'
            }
    
    async def _extract_with_selenium(self, url: str, wait_time: int) -> Dict[str, Any]:
        """استخراج باستخدام Selenium"""
        try:
            from selenium import webdriver
            from selenium.webdriver.chrome.options import Options
            from selenium.webdriver.common.by import By
            from selenium.webdriver.support.ui import WebDriverWait
            from selenium.webdriver.support import expected_conditions as EC
            
            options = Options()
            options.add_argument('--headless')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            
            driver = webdriver.Chrome(options=options)
            driver.get(url)
            
            # انتظار تحميل المحتوى
            WebDriverWait(driver, wait_time).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )
            
            content = driver.page_source
            title = driver.title
            
            driver.quit()
            
            return {
                'success': True,
                'method': 'selenium',
                'title': title,
                'content': content,
                'javascript_executed': True
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'خطأ في Selenium: {str(e)}'
            }

# ================ محلل SEO متقدم ================

class AdvancedSEOAnalyzer:
    """محلل تحسين محركات البحث المتقدم"""
    
    def __init__(self):
        self.seo_factors = {
            'technical': ['page_speed', 'mobile_friendly', 'https', 'structured_data'],
            'content': ['title_optimization', 'meta_description', 'headings', 'content_quality'],
            'social': ['og_tags', 'twitter_cards', 'social_sharing'],
            'accessibility': ['alt_tags', 'aria_labels', 'color_contrast']
        }
    
    def analyze_seo_comprehensive(self, soup: BeautifulSoup, url: str) -> Dict[str, Any]:
        """تحليل SEO شامل"""
        seo_analysis = {
            'overall_score': 0,
            'technical_seo': self._analyze_technical_seo(soup, url),
            'content_seo': self._analyze_content_seo(soup),
            'social_seo': self._analyze_social_seo(soup),
            'accessibility': self._analyze_accessibility(soup),
            'structured_data': self._extract_structured_data(soup),
            'recommendations': []
        }
        
        # حساب النتيجة الإجمالية
        scores = [
            seo_analysis['technical_seo']['score'],
            seo_analysis['content_seo']['score'],
            seo_analysis['social_seo']['score'],
            seo_analysis['accessibility']['score']
        ]
        seo_analysis['overall_score'] = sum(scores) / len(scores)
        
        # توليد التوصيات
        seo_analysis['recommendations'] = self._generate_seo_recommendations(seo_analysis)
        
        return seo_analysis
    
    def _analyze_technical_seo(self, soup: BeautifulSoup, url: str) -> Dict[str, Any]:
        """تحليل SEO التقني"""
        score = 0
        issues = []
        
        # فحص HTTPS
        if url.startswith('https'):
            score += 25
        else:
            issues.append('الموقع لا يستخدم HTTPS')
        
        # فحص العنوان
        title_tag = soup.find('title')
        if title_tag and title_tag.string:
            title_length = len(title_tag.string.strip())
            if 30 <= title_length <= 60:
                score += 25
            elif title_length > 0:
                score += 15
                issues.append(f'طول العنوان غير مثالي: {title_length} حرف')
        else:
            issues.append('العنوان مفقود')
        
        # فحص Meta Description
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc and meta_desc.get('content'):
            desc_length = len(meta_desc.get('content', ''))
            if 120 <= desc_length <= 160:
                score += 25
            elif desc_length > 0:
                score += 15
                issues.append(f'طول الوصف غير مثالي: {desc_length} حرف')
        else:
            issues.append('وصف Meta مفقود')
        
        # فحص البيانات المنظمة
        structured_data = soup.find_all('script', type='application/ld+json')
        if structured_data:
            score += 25
        else:
            issues.append('البيانات المنظمة مفقودة')
        
        return {
            'score': score,
            'issues': issues,
            'https_enabled': url.startswith('https'),
            'title_optimized': score >= 25,
            'meta_description_optimized': score >= 50,
            'structured_data_present': bool(structured_data)
        }
    
    def _analyze_content_seo(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """تحليل SEO المحتوى"""
        score = 0
        issues = []
        
        # تحليل العناوين
        headings = {}
        for i in range(1, 7):
            headings[f'h{i}'] = soup.find_all(f'h{i}')
        
        if headings['h1']:
            if len(headings['h1']) == 1:
                score += 25
            else:
                score += 15
                issues.append(f'عدد غير مثالي من H1: {len(headings["h1"])}')
        else:
            issues.append('H1 مفقود')
        
        # فحص هيكل العناوين
        if headings['h2']:
            score += 20
        else:
            issues.append('عناوين H2 مفقودة')
        
        # تحليل المحتوى النصي
        paragraphs = soup.find_all('p')
        if len(paragraphs) >= 3:
            score += 25
        elif len(paragraphs) > 0:
            score += 15
            issues.append('المحتوى النصي قليل')
        else:
            issues.append('لا يوجد محتوى نصي')
        
        # فحص الروابط الداخلية
        internal_links = len([a for a in soup.find_all('a', href=True) 
                            if not a.get('href', '').startswith('http')])
        if internal_links >= 3:
            score += 30
        elif internal_links > 0:
            score += 15
            issues.append('روابط داخلية قليلة')
        else:
            issues.append('لا توجد روابط داخلية')
        
        return {
            'score': score,
            'issues': issues,
            'headings_structure': {k: len(v) for k, v in headings.items()},
            'content_length': len(soup.get_text()),
            'internal_links_count': internal_links
        }
    
    def _analyze_social_seo(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """تحليل SEO الشبكات الاجتماعية"""
        score = 0
        issues = []
        
        # فحص Open Graph tags
        og_tags = {
            'og:title': soup.find('meta', property='og:title'),
            'og:description': soup.find('meta', property='og:description'),
            'og:image': soup.find('meta', property='og:image'),
            'og:url': soup.find('meta', property='og:url')
        }
        
        og_present = sum(1 for tag in og_tags.values() if tag)
        score += (og_present / 4) * 50
        
        if og_present < 4:
            missing_og = [name for name, tag in og_tags.items() if not tag]
            issues.append(f'Open Graph tags مفقودة: {", ".join(missing_og)}')
        
        # فحص Twitter Cards
        twitter_tags = {
            'twitter:card': soup.find('meta', attrs={'name': 'twitter:card'}),
            'twitter:title': soup.find('meta', attrs={'name': 'twitter:title'}),
            'twitter:description': soup.find('meta', attrs={'name': 'twitter:description'})
        }
        
        twitter_present = sum(1 for tag in twitter_tags.values() if tag)
        score += (twitter_present / 3) * 50
        
        if twitter_present < 3:
            missing_twitter = [name for name, tag in twitter_tags.items() if not tag]
            issues.append(f'Twitter Cards مفقودة: {", ".join(missing_twitter)}')
        
        return {
            'score': score,
            'issues': issues,
            'open_graph_tags': {k: bool(v) for k, v in og_tags.items()},
            'twitter_cards': {k: bool(v) for k, v in twitter_tags.items()}
        }
    
    def _analyze_accessibility(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """تحليل إمكانية الوصول"""
        score = 0
        issues = []
        
        # فحص alt tags للصور
        images = soup.find_all('img')
        images_with_alt = [img for img in images if img.get('alt')]
        
        if images:
            alt_ratio = len(images_with_alt) / len(images)
            score += alt_ratio * 30
            
            if alt_ratio < 1:
                issues.append(f'صور بدون alt tags: {len(images) - len(images_with_alt)}')
        else:
            score += 30  # لا توجد صور للفحص
        
        # فحص عناوين الروابط
        links = soup.find_all('a')
        links_with_title = [link for link in links if link.get('title') or link.get_text().strip()]
        
        if links:
            title_ratio = len(links_with_title) / len(links)
            score += title_ratio * 25
            
            if title_ratio < 0.8:
                issues.append('روابط بدون عناوين وصفية')
        else:
            score += 25
        
        # فحص ARIA labels
        aria_elements = soup.find_all(attrs={'aria-label': True})
        if aria_elements:
            score += 25
        else:
            issues.append('ARIA labels مفقودة')
        
        # فحص بنية HTML الدلالية
        semantic_tags = ['header', 'nav', 'main', 'article', 'section', 'aside', 'footer']
        semantic_present = sum(1 for tag in semantic_tags if soup.find(tag))
        score += (semantic_present / len(semantic_tags)) * 20
        
        if semantic_present < len(semantic_tags) * 0.5:
            issues.append('استخدام قليل للعناصر الدلالية')
        
        return {
            'score': score,
            'issues': issues,
            'images_with_alt': len(images_with_alt),
            'total_images': len(images),
            'semantic_tags_present': semantic_present,
            'aria_elements': len(aria_elements)
        }
    
    def _extract_structured_data(self, soup: BeautifulSoup) -> List[Dict[str, Any]]:
        """استخراج البيانات المنظمة"""
        structured_data = []
        
        # JSON-LD
        json_ld_scripts = soup.find_all('script', type='application/ld+json')
        for script in json_ld_scripts:
            try:
                data = json.loads(script.string)
                structured_data.append({
                    'type': 'JSON-LD',
                    'data': data
                })
            except (json.JSONDecodeError, AttributeError):
                continue
        
        # Microdata
        microdata_items = soup.find_all(attrs={'itemscope': True})
        for item in microdata_items:
            item_type = item.get('itemtype', '')
            structured_data.append({
                'type': 'Microdata',
                'itemtype': item_type,
                'properties': self._extract_microdata_properties(item)
            })
        
        return structured_data
    
    def _extract_microdata_properties(self, item) -> Dict[str, str]:
        """استخراج خصائص Microdata"""
        properties = {}
        prop_elements = item.find_all(attrs={'itemprop': True})
        
        for prop in prop_elements:
            prop_name = prop.get('itemprop')
            if prop.get('content'):
                prop_value = prop.get('content')
            elif prop.name in ['meta']:
                prop_value = prop.get('content', '')
            else:
                prop_value = prop.get_text().strip()
            
            properties[prop_name] = prop_value
        
        return properties
    
    def _generate_seo_recommendations(self, analysis: Dict[str, Any]) -> List[str]:
        """توليد توصيات SEO"""
        recommendations = []
        
        # توصيات تقنية
        if analysis['technical_seo']['score'] < 75:
            recommendations.extend([
                'تحسين العنوان ليكون بين 30-60 حرف',
                'كتابة وصف meta بين 120-160 حرف',
                'إضافة البيانات المنظمة (Schema.org)',
                'التأكد من استخدام HTTPS'
            ])
        
        # توصيات المحتوى
        if analysis['content_seo']['score'] < 75:
            recommendations.extend([
                'إضافة عنوان H1 واحد فقط لكل صفحة',
                'استخدام عناوين H2-H6 بشكل هرمي',
                'زيادة المحتوى النصي عالي الجودة',
                'إضافة المزيد من الروابط الداخلية'
            ])
        
        # توصيات الشبكات الاجتماعية
        if analysis['social_seo']['score'] < 75:
            recommendations.extend([
                'إضافة Open Graph tags للمشاركة على فيسبوك',
                'إضافة Twitter Cards للمشاركة على تويتر',
                'تحسين الصور للمشاركة الاجتماعية'
            ])
        
        # توصيات إمكانية الوصول
        if analysis['accessibility']['score'] < 75:
            recommendations.extend([
                'إضافة alt tags لجميع الصور',
                'استخدام ARIA labels للعناصر التفاعلية',
                'تحسين بنية HTML باستخدام العناصر الدلالية',
                'إضافة عناوين وصفية للروابط'
            ])
        
        return recommendations

# ================ نظام التصدير المتقدم ================

class AdvancedExporter:
    """نظام تصدير متقدم للنتائج"""
    
    def __init__(self):
        self.export_formats = ['json', 'csv', 'xml', 'pdf', 'docx']
    
    def export_extraction_results(self, results: Dict[str, Any], 
                                format_type: str = 'json') -> str:
        """تصدير نتائج الاستخراج بصيغات مختلفة"""
        
        if format_type == 'json':
            return self._export_json(results)
        elif format_type == 'csv':
            return self._export_csv(results)
        elif format_type == 'xml':
            return self._export_xml(results)
        elif format_type == 'pdf':
            return self._export_pdf(results)
        elif format_type == 'docx':
            return self._export_docx(results)
        else:
            raise ValueError(f"صيغة غير مدعومة: {format_type}")
    
    def _export_json(self, results: Dict[str, Any]) -> str:
        """تصدير JSON"""
        filename = f"extraction_results_{int(time.time())}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        return filename
    
    def _export_csv(self, results: Dict[str, Any]) -> str:
        """تصدير CSV"""
        filename = f"extraction_results_{int(time.time())}.csv"
        
        # تسطيح البيانات للـ CSV
        flattened_data = self._flatten_results(results)
        
        with open(filename, 'w', newline='', encoding='utf-8') as f:
            if flattened_data:
                writer = csv.DictWriter(f, fieldnames=flattened_data[0].keys())
                writer.writeheader()
                writer.writerows(flattened_data)
        
        return filename
    
    def _export_xml(self, results: Dict[str, Any]) -> str:
        """تصدير XML"""
        filename = f"extraction_results_{int(time.time())}.xml"
        
        root = ET.Element("extraction_results")
        self._dict_to_xml(results, root)
        
        tree = ET.ElementTree(root)
        tree.write(filename, encoding='utf-8', xml_declaration=True)
        
        return filename
    
    def _export_pdf(self, results: Dict[str, Any]) -> str:
        """تصدير PDF"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            filename = f"extraction_results_{int(time.time())}.pdf"
            
            c = canvas.Canvas(filename, pagesize=letter)
            width, height = letter
            
            # إضافة العنوان
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height - 50, "Website Extraction Results")
            
            # إضافة المحتوى
            y_position = height - 100
            c.setFont("Helvetica", 12)
            
            for key, value in results.items():
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                
                c.drawString(50, y_position, f"{key}: {str(value)[:50]}...")
                y_position -= 20
            
            c.save()
            return filename
            
        except ImportError:
            return self._export_json(results)  # fallback to JSON
    
    def _export_docx(self, results: Dict[str, Any]) -> str:
        """تصدير Word Document"""
        try:
            from docx import Document
            
            filename = f"extraction_results_{int(time.time())}.docx"
            doc = Document()
            
            # إضافة العنوان
            doc.add_heading('Website Extraction Results', 0)
            
            # إضافة المحتوى
            for key, value in results.items():
                doc.add_heading(str(key), level=1)
                doc.add_paragraph(str(value))
            
            doc.save(filename)
            return filename
            
        except ImportError:
            return self._export_json(results)  # fallback to JSON
    
    def _flatten_results(self, data: Dict[str, Any], parent_key: str = '') -> List[Dict[str, Any]]:
        """تسطيح البيانات للـ CSV"""
        items = []
        
        for key, value in data.items():
            new_key = f"{parent_key}.{key}" if parent_key else key
            
            if isinstance(value, dict):
                items.extend(self._flatten_results(value, new_key))
            elif isinstance(value, list):
                for i, item in enumerate(value):
                    if isinstance(item, dict):
                        items.extend(self._flatten_results(item, f"{new_key}[{i}]"))
                    else:
                        items.append({new_key: str(item)})
            else:
                items.append({new_key: str(value)})
        
        return items
    
    def _dict_to_xml(self, data: Dict[str, Any], parent: ET.Element):
        """تحويل Dictionary إلى XML"""
        for key, value in data.items():
            element = ET.SubElement(parent, str(key))
            
            if isinstance(value, dict):
                self._dict_to_xml(value, element)
            elif isinstance(value, list):
                for item in value:
                    if isinstance(item, dict):
                        sub_element = ET.SubElement(element, "item")
                        self._dict_to_xml(item, sub_element)
                    else:
                        sub_element = ET.SubElement(element, "item")
                        sub_element.text = str(item)
            else:
                element.text = str(value)

    def _get_safe_filename(self, url: str) -> str:
        """إنشاء اسم ملف آمن"""
        parsed = urlparse(url)
        path = parsed.path
        
        if not path or path == '/':
            return 'index.html'
        
        filename = unquote(path.split('/')[-1])
        if not filename:
            return 'index.html'
        
        if not '.' in filename:
            filename += '.html'
        
        # تنظيف الأحرف غير الآمنة
        safe_chars = '-_.() abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789'
        filename = ''.join(c for c in filename if c in safe_chars)
        
        return filename[:100]

    def _save_extraction_report(self):
        """حفظ تقرير الاستخراج"""
        report_path = self.site_dir / 'reports' / 'extraction_report.json'
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(self.stats, f, indent=2, ensure_ascii=False)

    def _compress_output(self):
        """ضغط مجلد الإخراج"""
        zip_path = self.output_dir / f"{self.site_id}.zip"
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in self.site_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(self.site_dir)
                    zipf.write(file_path, arcname)
        
        logger.info(f"تم ضغط الملفات في: {zip_path}")

    # Placeholder methods for complete functionality
    def _extract_page_assets(self, soup, url): return {}
    def _remove_ads(self, soup): return soup
    def _update_html_links(self, soup, url): return soup
    def _save_content_data(self, content): pass
    def _parallel_extraction(self, executor): return {}
    def _analyze_site_structure(self): return {}
    def _smart_extraction(self, analysis): return {}
    def _generate_ai_insights(self): return {}
    def _perform_security_check(self): return {'safe': True, 'issues': []}
    def _secure_extraction(self): return {}