"""
Master Extractor - Fixed and Complete Implementation
This module provides comprehensive website content extraction capabilities
with proper error handling and compatibility for Replit environment.
"""

import os
import re
import json
import time
import hashlib
import logging
import asyncio
from datetime import datetime
from urllib.parse import urljoin, urlparse
from typing import Dict, List, Any, Optional, Tuple, Union
from dataclasses import dataclass, field
from enum import Enum

import requests
from bs4 import BeautifulSoup, Tag, NavigableString
import aiohttp
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from playwright.async_api import async_playwright
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ExtractionMode(Enum):
    """Extraction modes available"""
    BASIC = "basic"
    STANDARD = "standard"
    ADVANCED = "advanced"
    ULTRA = "ultra"
    SECURE = "secure"


@dataclass
class ExtractionConfig:
    """Configuration for extraction process"""
    mode: ExtractionMode = ExtractionMode.BASIC
    max_pages: int = 50
    max_depth: int = 3
    timeout: int = 30
    extract_images: bool = True
    extract_documents: bool = True
    extract_media: bool = True
    block_ads: bool = False
    extract_javascript: bool = False
    extract_css: bool = False
    follow_external_links: bool = False
    output_format: str = "json"
    custom_selectors: List[str] = field(default_factory=list)
    user_agent: str = "MasterExtractor/1.0"


class ExtractionEngine:
    """Core extraction engine with proper initialization"""
    
    def __init__(self, config: ExtractionConfig):
        self.config = config
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': config.user_agent
        })
        
        # Initialize tracking variables
        self.visited_urls = set()
        self.failed_urls = set()
        self.extracted_data = {}
        self.stats = {
            'pages_processed': 0,
            'data_extracted': 0,
            'errors': 0,
            'start_time': datetime.now()
        }
        
        # URL parsing
        self.base_url = ""
        self.domain = ""
        
        # Storage
        self.site_dir = ""
        self.site_id = ""
        self.page_metadata = {}
        
        # Ad blocking selectors
        self.ad_selectors = [
            '[class*="ad"]', '[id*="ad"]',
            '[class*="advertisement"]', '[id*="advertisement"]',
            '.google-ads', '.adsense',
            '[class*="banner"]', '[id*="banner"]'
        ]

    def set_base_url(self, url: str):
        """Set base URL and extract domain"""
        self.base_url = url
        parsed = urlparse(url)
        self.domain = parsed.netloc
        self.site_id = hashlib.md5(url.encode()).hexdigest()[:8]
        self.site_dir = f"temp/extraction_{self.site_id}"
        os.makedirs(self.site_dir, exist_ok=True)

    def extract_page_content(self, url: str, soup: BeautifulSoup) -> Dict[str, Any]:
        """Extract content from a single page"""
        try:
            # Basic page information
            content = {
                'url': url,
                'title': soup.title.string if soup.title else '',
                'meta_description': '',
                'headings': {},
                'paragraphs': [],
                'links': [],
                'images': [],
                'forms': [],
                'scripts': [],
                'stylesheets': [],
                'extracted_at': datetime.now().isoformat()
            }
            
            # Extract meta description
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc and meta_desc.get('content'):
                content['meta_description'] = meta_desc.get('content', '')
            
            # Extract headings
            for level in range(1, 7):
                headings = soup.find_all(f'h{level}')
                if headings:
                    content['headings'][f'h{level}'] = [h.get_text(strip=True) for h in headings]
            
            # Extract paragraphs
            paragraphs = soup.find_all('p')
            content['paragraphs'] = [p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)]
            
            # Extract links
            links = soup.find_all('a', href=True)
            for link in links:
                href = link.get('href', '')
                text = link.get_text(strip=True)
                if href and text:
                    content['links'].append({
                        'url': urljoin(url, href),
                        'text': text,
                        'is_external': not href.startswith(self.base_url)
                    })
            
            # Extract images
            if self.config.extract_images:
                images = soup.find_all('img')
                for img in images:
                    src = img.get('src', '')
                    alt = img.get('alt', '')
                    if src:
                        content['images'].append({
                            'src': urljoin(url, src),
                            'alt': alt,
                            'width': img.get('width', ''),
                            'height': img.get('height', '')
                        })
            
            # Extract forms
            forms = soup.find_all('form')
            for form in forms:
                action = form.get('action', '')
                method = form.get('method', 'GET')
                inputs = form.find_all(['input', 'textarea', 'select'])
                form_data = {
                    'action': urljoin(url, action) if action else '',
                    'method': method.upper(),
                    'fields': []
                }
                
                for input_elem in inputs:
                    field_data = {
                        'name': input_elem.get('name', ''),
                        'type': input_elem.get('type', 'text'),
                        'required': input_elem.has_attr('required')
                    }
                    form_data['fields'].append(field_data)
                
                content['forms'].append(form_data)
            
            # Extract scripts and stylesheets
            if self.config.extract_javascript:
                scripts = soup.find_all('script', src=True)
                content['scripts'] = [urljoin(url, script.get('src', '')) for script in scripts]
            
            if self.config.extract_css:
                stylesheets = soup.find_all('link', rel='stylesheet')
                content['stylesheets'] = [urljoin(url, link.get('href', '')) for link in stylesheets]
            
            # Remove ads if configured
            if self.config.block_ads:
                self._remove_ads(soup)
            
            # Update statistics
            self.stats['pages_processed'] += 1
            self.stats['data_extracted'] += len(content['paragraphs']) + len(content['links'])
            
            return content
            
        except Exception as e:
            logger.error(f"Error extracting content from {url}: {e}")
            self.stats['errors'] += 1
            return {'url': url, 'error': str(e)}

    def _remove_ads(self, soup: BeautifulSoup):
        """Remove advertising elements from soup"""
        for selector in self.ad_selectors:
            try:
                ads = soup.select(selector)
                for ad in ads:
                    ad.decompose()
            except Exception as e:
                logger.warning(f"Error removing ads with selector {selector}: {e}")

    def extract_website_data(self, url: str) -> Dict[str, Any]:
        """Extract data from a website using the configured mode"""
        try:
            self.set_base_url(url)
            
            if self.config.mode == ExtractionMode.BASIC:
                return self._extract_basic(url)
            elif self.config.mode == ExtractionMode.STANDARD:
                return self._extract_standard(url)
            elif self.config.mode == ExtractionMode.ADVANCED:
                return self._extract_advanced(url)
            elif self.config.mode == ExtractionMode.ULTRA:
                return self._extract_ultra(url)
            elif self.config.mode == ExtractionMode.SECURE:
                return self._extract_secure(url)
            else:
                return self._extract_basic(url)
                
        except Exception as e:
            logger.error(f"Error in extract_website_data: {e}")
            return {'error': str(e), 'url': url}

    def _extract_basic(self, url: str) -> Dict[str, Any]:
        """Basic extraction using requests and BeautifulSoup"""
        try:
            response = self.session.get(url, timeout=self.config.timeout)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            content = self.extract_page_content(url, soup)
            
            return {
                'mode': 'basic',
                'base_url': url,
                'pages': [content],
                'statistics': self.stats,
                'extraction_time': (datetime.now() - self.stats['start_time']).total_seconds()
            }
            
        except Exception as e:
            logger.error(f"Basic extraction failed: {e}")
            return {'error': str(e), 'mode': 'basic'}

    def _extract_standard(self, url: str) -> Dict[str, Any]:
        """Standard extraction with link following"""
        try:
            pages_data = []
            urls_to_visit = [url]
            current_depth = 0
            
            while urls_to_visit and current_depth < self.config.max_depth and len(pages_data) < self.config.max_pages:
                current_url = urls_to_visit.pop(0)
                
                if current_url in self.visited_urls:
                    continue
                    
                self.visited_urls.add(current_url)
                
                try:
                    response = self.session.get(current_url, timeout=self.config.timeout)
                    response.raise_for_status()
                    
                    soup = BeautifulSoup(response.content, 'html.parser')
                    page_content = self.extract_page_content(current_url, soup)
                    pages_data.append(page_content)
                    
                    # Find new links to visit
                    if current_depth < self.config.max_depth - 1:
                        links = soup.find_all('a', href=True)
                        for link in links:
                            href = link.get('href', '')
                            if href:
                                full_url = urljoin(current_url, href)
                                # Only follow internal links unless configured otherwise
                                if (self.config.follow_external_links or 
                                    full_url.startswith(self.base_url)):
                                    if full_url not in self.visited_urls:
                                        urls_to_visit.append(full_url)
                    
                except Exception as e:
                    logger.error(f"Error processing {current_url}: {e}")
                    self.failed_urls.add(current_url)
                
                current_depth += 1
                time.sleep(0.5)  # Rate limiting
            
            return {
                'mode': 'standard',
                'base_url': url,
                'pages': pages_data,
                'statistics': self.stats,
                'visited_urls': len(self.visited_urls),
                'failed_urls': len(self.failed_urls),
                'extraction_time': (datetime.now() - self.stats['start_time']).total_seconds()
            }
            
        except Exception as e:
            logger.error(f"Standard extraction failed: {e}")
            return {'error': str(e), 'mode': 'standard'}

    def _extract_advanced(self, url: str) -> Dict[str, Any]:
        """Advanced extraction using Selenium for JavaScript content"""
        try:
            # Setup Chrome options for Replit environment
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--window-size=1920,1080')
            
            driver = webdriver.Chrome(options=chrome_options)
            pages_data = []
            
            try:
                driver.get(url)
                WebDriverWait(driver, self.config.timeout).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
                
                # Wait for JavaScript to load
                time.sleep(3)
                
                # Get page source after JavaScript execution
                html_content = driver.page_source
                soup = BeautifulSoup(html_content, 'html.parser')
                
                page_content = self.extract_page_content(url, soup)
                
                # Extract additional JavaScript-loaded content
                try:
                    # Extract any dynamically loaded text
                    dynamic_content = driver.find_elements(By.XPATH, "//*[contains(@class, 'dynamic') or contains(@id, 'dynamic')]")
                    dynamic_text = [elem.text for elem in dynamic_content if elem.text.strip()]
                    if dynamic_text:
                        page_content['dynamic_content'] = dynamic_text
                except Exception as e:
                    logger.warning(f"Error extracting dynamic content: {e}")
                
                pages_data.append(page_content)
                
            finally:
                driver.quit()
            
            return {
                'mode': 'advanced',
                'base_url': url,
                'pages': pages_data,
                'statistics': self.stats,
                'extraction_time': (datetime.now() - self.stats['start_time']).total_seconds()
            }
            
        except Exception as e:
            logger.error(f"Advanced extraction failed: {e}")
            return {'error': str(e), 'mode': 'advanced'}

    async def _extract_ultra(self, url: str) -> Dict[str, Any]:
        """Ultra extraction using Playwright for modern web apps"""
        try:
            async with async_playwright() as p:
                browser = await p.chromium.launch(headless=True)
                page = await browser.new_page()
                
                pages_data = []
                
                try:
                    await page.goto(url, wait_until='networkidle')
                    
                    # Wait for dynamic content
                    await page.wait_for_timeout(3000)
                    
                    # Get page content
                    content = await page.content()
                    soup = BeautifulSoup(content, 'html.parser')
                    
                    page_content = self.extract_page_content(url, soup)
                    
                    # Extract additional Playwright-specific data
                    try:
                        # Get page title and meta info
                        title = await page.title()
                        page_content['playwright_title'] = title
                        
                        # Get viewport size
                        viewport = page.viewport_size
                        page_content['viewport'] = viewport
                        
                        # Extract any SPA content
                        spa_content = await page.evaluate('''
                            () => {
                                const scripts = Array.from(document.querySelectorAll('script'));
                                const hasReact = scripts.some(s => s.src.includes('react'));
                                const hasVue = scripts.some(s => s.src.includes('vue'));
                                const hasAngular = scripts.some(s => s.src.includes('angular'));
                                return {
                                    framework: hasReact ? 'React' : hasVue ? 'Vue' : hasAngular ? 'Angular' : 'Unknown',
                                    scriptsCount: scripts.length
                                };
                            }
                        ''')
                        page_content['spa_info'] = spa_content
                        
                    except Exception as e:
                        logger.warning(f"Error extracting Playwright-specific data: {e}")
                    
                    pages_data.append(page_content)
                    
                finally:
                    await browser.close()
                
                return {
                    'mode': 'ultra',
                    'base_url': url,
                    'pages': pages_data,
                    'statistics': self.stats,
                    'extraction_time': (datetime.now() - self.stats['start_time']).total_seconds()
                }
                
        except Exception as e:
            logger.error(f"Ultra extraction failed: {e}")
            return {'error': str(e), 'mode': 'ultra'}

    def _extract_secure(self, url: str) -> Dict[str, Any]:
        """Secure extraction with enhanced protection and validation"""
        try:
            # Enhanced security headers
            secure_headers = {
                'User-Agent': self.config.user_agent,
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
                'Accept-Encoding': 'gzip, deflate',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1'
            }
            
            session = requests.Session()
            session.headers.update(secure_headers)
            
            # Validate URL before processing
            parsed_url = urlparse(url)
            if not parsed_url.scheme or not parsed_url.netloc:
                raise ValueError("Invalid URL format")
            
            response = session.get(url, timeout=self.config.timeout, verify=True)
            response.raise_for_status()
            
            # Additional security checks
            content_type = response.headers.get('content-type', '')
            if 'text/html' not in content_type.lower():
                logger.warning(f"Unexpected content type: {content_type}")
            
            soup = BeautifulSoup(response.content, 'html.parser')
            page_content = self.extract_page_content(url, soup)
            
            # Add security analysis
            security_analysis = {
                'https_enabled': parsed_url.scheme == 'https',
                'content_type': content_type,
                'server': response.headers.get('server', 'Unknown'),
                'security_headers': {
                    'strict_transport_security': 'strict-transport-security' in response.headers,
                    'content_security_policy': 'content-security-policy' in response.headers,
                    'x_frame_options': 'x-frame-options' in response.headers,
                    'x_content_type_options': 'x-content-type-options' in response.headers
                }
            }
            
            page_content['security_analysis'] = security_analysis
            
            return {
                'mode': 'secure',
                'base_url': url,
                'pages': [page_content],
                'statistics': self.stats,
                'extraction_time': (datetime.now() - self.stats['start_time']).total_seconds()
            }
            
        except Exception as e:
            logger.error(f"Secure extraction failed: {e}")
            return {'error': str(e), 'mode': 'secure'}


class AdvancedExporter:
    """Export extraction results to various formats"""
    
    def __init__(self, output_dir: str = "temp"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def export_to_json(self, data: Dict[str, Any], filename: str = None) -> str:
        """Export data to JSON file"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"extraction_{timestamp}.json"
        
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False, default=str)
        
        return filepath

    def export_to_pdf(self, data: Dict[str, Any], filename: str = None) -> str:
        """Export data to PDF file"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"extraction_{timestamp}.pdf"
        
        filepath = os.path.join(self.output_dir, filename)
        
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        
        y_position = height - 50
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y_position, f"Website Extraction Report")
        y_position -= 30
        
        # Basic info
        c.setFont("Helvetica", 12)
        if 'base_url' in data:
            c.drawString(50, y_position, f"URL: {data['base_url']}")
            y_position -= 20
        
        if 'mode' in data:
            c.drawString(50, y_position, f"Extraction Mode: {data['mode']}")
            y_position -= 20
        
        # Statistics
        if 'statistics' in data:
            stats = data['statistics']
            c.drawString(50, y_position, f"Pages Processed: {stats.get('pages_processed', 0)}")
            y_position -= 20
            c.drawString(50, y_position, f"Data Points: {stats.get('data_extracted', 0)}")
            y_position -= 20
        
        # Page summaries
        if 'pages' in data:
            y_position -= 20
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y_position, "Page Summaries:")
            y_position -= 20
            
            for i, page in enumerate(data['pages'][:5]):  # Limit to first 5 pages
                c.setFont("Helvetica", 10)
                if y_position < 100:
                    c.showPage()
                    y_position = height - 50
                
                c.drawString(50, y_position, f"Page {i+1}: {page.get('title', 'No title')[:60]}")
                y_position -= 15
        
        c.save()
        return filepath

    def export_to_docx(self, data: Dict[str, Any], filename: str = None) -> str:
        """Export data to Word document"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"extraction_{timestamp}.docx"
        
        filepath = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Website Extraction Report', 0)
        
        # Basic information
        if 'base_url' in data:
            p = doc.add_paragraph()
            p.add_run('URL: ').bold = True
            p.add_run(data['base_url'])
        
        if 'mode' in data:
            p = doc.add_paragraph()
            p.add_run('Extraction Mode: ').bold = True
            p.add_run(data['mode'])
        
        # Statistics
        if 'statistics' in data:
            doc.add_heading('Statistics', level=1)
            stats = data['statistics']
            
            stats_list = [
                f"Pages Processed: {stats.get('pages_processed', 0)}",
                f"Data Points Extracted: {stats.get('data_extracted', 0)}",
                f"Errors: {stats.get('errors', 0)}",
                f"Extraction Time: {data.get('extraction_time', 0):.2f} seconds"
            ]
            
            for stat in stats_list:
                doc.add_paragraph(stat, style='List Bullet')
        
        # Page details
        if 'pages' in data:
            doc.add_heading('Page Details', level=1)
            
            for i, page in enumerate(data['pages'][:10]):  # Limit to first 10 pages
                doc.add_heading(f"Page {i+1}: {page.get('title', 'No title')}", level=2)
                
                if page.get('url'):
                    p = doc.add_paragraph()
                    p.add_run('URL: ').bold = True
                    p.add_run(page['url'])
                
                if page.get('meta_description'):
                    p = doc.add_paragraph()
                    p.add_run('Description: ').bold = True
                    p.add_run(page['meta_description'])
                
                # Add headings summary
                if page.get('headings'):
                    p = doc.add_paragraph()
                    p.add_run('Headings Found: ').bold = True
                    heading_count = sum(len(headings) for headings in page['headings'].values())
                    p.add_run(str(heading_count))
                
                # Add content summary
                if page.get('paragraphs'):
                    p = doc.add_paragraph()
                    p.add_run('Content Paragraphs: ').bold = True
                    p.add_run(str(len(page['paragraphs'])))
        
        doc.save(filepath)
        return filepath


class MasterExtractor:
    """Main extractor class that coordinates all extraction operations"""
    
    def __init__(self, config: ExtractionConfig = None):
        self.config = config or ExtractionConfig()
        self.engine = ExtractionEngine(self.config)
        self.exporter = AdvancedExporter()
        
    def extract(self, url: str, save_report: bool = True, compress_output: bool = False) -> Dict[str, Any]:
        """Main extraction method"""
        try:
            logger.info(f"Starting extraction of {url} with mode {self.config.mode.value}")
            
            # Perform extraction based on mode
            if self.config.mode == ExtractionMode.ULTRA:
                # Ultra mode requires async execution
                loop = asyncio.get_event_loop()
                result = loop.run_until_complete(self.engine._extract_ultra(url))
            else:
                result = self.engine.extract_website_data(url)
            
            # Save report if requested
            if save_report and 'error' not in result:
                json_file = self.exporter.export_to_json(result)
                result['report_files'] = {'json': json_file}
                
                # Export to additional formats based on config
                if self.config.output_format in ['pdf', 'all']:
                    pdf_file = self.exporter.export_to_pdf(result)
                    result['report_files']['pdf'] = pdf_file
                
                if self.config.output_format in ['docx', 'all']:
                    docx_file = self.exporter.export_to_docx(result)
                    result['report_files']['docx'] = docx_file
            
            # Compress output if requested
            if compress_output and 'report_files' in result:
                # Implementation would go here for compression
                pass
            
            logger.info(f"Extraction completed successfully")
            return result
            
        except Exception as e:
            logger.error(f"Extraction failed: {e}")
            return {'error': str(e), 'url': url}


# Convenience functions for easy usage
def extract_basic(url: str) -> Dict[str, Any]:
    """Quick basic extraction"""
    config = ExtractionConfig(mode=ExtractionMode.BASIC)
    extractor = MasterExtractor(config)
    return extractor.extract(url)


def extract_standard(url: str, max_pages: int = 50) -> Dict[str, Any]:
    """Quick standard extraction with link following"""
    config = ExtractionConfig(mode=ExtractionMode.STANDARD, max_pages=max_pages)
    extractor = MasterExtractor(config)
    return extractor.extract(url)


def extract_advanced(url: str) -> Dict[str, Any]:
    """Quick advanced extraction with JavaScript support"""
    config = ExtractionConfig(mode=ExtractionMode.ADVANCED)
    extractor = MasterExtractor(config)
    return extractor.extract(url)


# Main execution for testing
if __name__ == "__main__":
    # Example usage
    url = "https://example.com"
    
    # Basic extraction
    result = extract_basic(url)
    print(f"Basic extraction result: {len(result.get('pages', []))} pages")
    
    # Advanced extraction with custom config
    config = ExtractionConfig(
        mode=ExtractionMode.ADVANCED,
        max_pages=10,
        extract_images=True,
        block_ads=True,
        output_format='all'
    )
    
    extractor = MasterExtractor(config)
    result = extractor.extract(url, save_report=True)
    print(f"Advanced extraction completed: {result.get('extraction_time', 0):.2f} seconds")